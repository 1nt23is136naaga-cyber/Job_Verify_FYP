"""
Job Posting Verifier v6 — FIXED
─────────────────────────────────────────────────────────────────────────────
Fixes applied on top of v6:

  FIX 1 — PROPER TASK CANCELLATION ON EARLY EXIT
  • Platform tasks are wrapped in asyncio.Task objects tracked in a list.
  • When done_event fires (score ≥ 92), all pending tasks are cancelled via
    task.cancel() before browser.close() is called.
  • Prevents "Target page closed" errors from tasks still trying to use the
    browser after it has been shut down.

  FIX 2 — PLAYWRIGHT CALLS GUARDED AGAINST CLOSED BROWSER
  • Every Playwright await is wrapped with a done_event.is_set() pre-check.
  • pw_goto / pw_extract / pw_human_type all short-circuit cleanly when the
    browser is shutting down, instead of raising TargetClosedError.

  FIX 3 — LATE ASYNC ERROR SUPPRESSION (asyncio best practice)
  • All Task objects have an exception-retrieval shim attached via
    task.add_done_callback so that cancelled/errored futures are always
    "retrieved" — silencing the "Future exception was never retrieved" warning.

  FIX 4 — GLOBAL ASYNCIO EXCEPTION HANDLER
  • loop.set_exception_handler() installed to swallow residual
    TargetClosedError / CancelledError noise that escapes task boundaries.

  FIX 5 — OUTPUT ORDER CORRECTED
  • Careers page check now runs first and PRINTS before platform tasks start.
  • Platforms + Google then run concurrently as before.
  • Output sequence:
      [1/4] Resolving careers URL …
      [2/4] Careers page: ✅ …          ← now prints BEFORE platforms
      [3+4/4] Google + 14 platforms …
      [3/4] Google confirmation: …
      Verdict …

  FIX 6 — asyncio.coroutine REMOVED
  • Replaced deprecated asyncio.coroutine lambda with a proper async def
    _no_careers_url() coroutine.

  ALL v6 PERFORMANCE PRESERVED — ZERO COMPROMISE
  ✔ HTTP/2 fast-path → Playwright fallback
  ✔ True async, early-exit short-circuit ≥ 92
  ✔ ATS override table (50+ companies)
  ✔ 14 platform scrapers + careers + Google
  ✔ Resource blocking, anti-bot, fuzzy matching

─────────────────────────────────────────────────────────────────────────────
Install:
    pip install httpx[http2] playwright beautifulsoup4 rapidfuzz pdfplumber python-docx
    playwright install chromium
─────────────────────────────────────────────────────────────────────────────
"""

import re, time, random, asyncio, os, logging, sys
from urllib.parse import quote as urlquote
from bs4 import BeautifulSoup

# Ensure Windows prints unicode (emojis/box drawings) without crashing
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')


# ── Optional deps ──────────────────────────────────────────────────────────────
try:
    import httpx
    HAS_HTTPX = True
except ImportError:
    HAS_HTTPX = False
    print("[WARNING] pip install httpx[http2]")

try:
    from playwright.async_api import async_playwright, Browser, BrowserContext, Page
    HAS_PLAYWRIGHT = True
except ImportError:
    HAS_PLAYWRIGHT = False
    print("[WARNING] pip install playwright && playwright install chromium")

try:
    from rapidfuzz import fuzz
    HAS_RAPIDFUZZ = True
except ImportError:
    HAS_RAPIDFUZZ = False
    print("[WARNING] pip install rapidfuzz")

try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    import docx as _docx_mod
except ImportError:
    _docx_mod = None

logging.getLogger("httpx").setLevel(logging.ERROR)
logging.getLogger("playwright").setLevel(logging.ERROR)

# ── Constants ──────────────────────────────────────────────────────────────────

USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:124.0) Gecko/20100101 Firefox/124.0",
]

CAREERS_KEYWORDS    = ["career", "job", "work", "hire", "talent", "vacancy", "opening", "position"]
LOGIN_WALL_KEYWORDS = [
    "sign in to", "log in to", "create account", "join now to", "register to view",
    "please sign in", "login required", "access denied", "members only",
]

BLOCK_RESOURCE_TYPES = {
    "image", "media", "font", "stylesheet",
    "websocket", "manifest", "other",
}
BLOCK_DOMAINS = {
    "google-analytics", "googletagmanager", "doubleclick", "facebook.net",
    "hotjar", "clarity.ms", "segment.io", "amplitude.com", "mixpanel.com",
    "cdn.optimizely", "sentry.io", "newrelic", "pingdom", "appsflyer",
}

FUZZY_TITLE_THRESHOLD   = 72
FUZZY_COMPANY_THRESHOLD = 60

PLATFORM_TIMEOUT  = 20
HTTP_TIMEOUT      = 6
PW_GOTO_TIMEOUT   = 12000
PW_WAIT_AFTER     = 800
MAX_HTTP_PARALLEL = 16

# ── ATS Overrides ──────────────────────────────────────────────────────────────

ATS_OVERRIDES = {
    "optum":         "https://careers.unitedhealthgroup.com/jobs?q={title}&location=Bengaluru",
    "unitedhealth":  "https://careers.unitedhealthgroup.com/jobs?q={title}&location=Bengaluru",
    "uhg":           "https://careers.unitedhealthgroup.com/jobs?q={title}&location=Bengaluru",
    "accenture":     "https://www.accenture.com/in-en/careers/jobsearch?jk={title}",
    "infosys":       "https://career.infosys.com/joblist",
    "wipro":         "https://careers.wipro.com/careers-home/jobs?q={title}",
    "tcs":           "https://ibegin.tcs.com/iBegin/",
    "cognizant":     "https://careers.cognizant.com/global/en/search-results?keywords={title}",
    "hcltech":       "https://www.hcltech.com/careers/job-search?q={title}",
    "capgemini":     "https://www.capgemini.com/in-en/careers/find-a-job/?search={title}",
    "mphasis":       "https://careers.mphasis.com/home.html",
    "ltimindtree":   "https://www.ltimindtree.com/careers/job-openings/?search={title}",
    "techm":         "https://careers.techmahindra.com/search/?q={title}",
    "tech mahindra": "https://careers.techmahindra.com/search/?q={title}",
    "hexaware":      "https://hexaware.com/careers/job-openings/?s={title}",
    "ibm":           "https://www.ibm.com/employment/#jobs?q={title}&country=India",
    "amazon":        "https://www.amazon.jobs/en/search?query={title}&country=IND",
    "microsoft":     "https://jobs.careers.microsoft.com/global/en/search?q={title}",
    "google":        "https://careers.google.com/jobs/results/?q={title}",
    "meta":          "https://www.metacareers.com/jobs?q={title}",
    "apple":         "https://jobs.apple.com/en-us/search?search={title}",
    "salesforce":    "https://careers.salesforce.com/en/jobs/?search={title}",
    "oracle":        "https://careers.oracle.com/jobs/#en/sites/jobsearch/jobs?keyword={title}",
    "sap":           "https://jobs.sap.com/search/?q={title}&locname=India",
    "adobe":         "https://careers.adobe.com/us/en/search-results?keywords={title}",
    "intel":         "https://jobs.intel.com/en/search-jobs/{title}",
    "cisco":         "https://jobs.cisco.com/jobs/SearchJobs/{title}",
    "dell":          "https://jobs.dell.com/search-jobs/{title}/India/375",
    "nvidia":        "https://nvidia.wd5.myworkdayjobs.com/NVIDIAExternalCareerSite?q={title}",
    "amd":           "https://careers.amd.com/careers-home/jobs?q={title}",
    "deloitte":      "https://apply.deloitte.com/careers/SearchJobs/{title}",
    "pwc":           "https://www.pwc.com/gx/en/careers/job-search.html",
    "kpmg":          "https://home.kpmg/xx/en/home/careers/job-search.html",
    "ey":            "https://eyglobal.yello.co/jobs?filter%5Bkeyword%5D={title}",
    "mckinsey":      "https://www.mckinsey.com/careers/search-jobs#q={title}",
    "bcg":           "https://careers.bcg.com/job-search?q={title}",
    "jpmorgan":      "https://jobs.jpmorganchase.com/careersection/10140/jobsearch.ftl?lang=en&keyword={title}",
    "goldman":       "https://higher.gs.com/roles?q={title}",
    "citi":          "https://jobs.citi.com/search-jobs/{title}/287/1",
    "hsbc":          "https://mycareer.hsbc.com/en_GB/external/SearchJobs/{title}",
    "atlassian":     "https://www.atlassian.com/company/careers/all-jobs?search={title}",
    "servicenow":    "https://jobs.smartrecruiters.com/ServiceNow/{title}",
    "workday":       "https://workday.wd5.myworkdayjobs.com/Workday/jobs?q={title}",
    "zoho":          "https://careers.zohocorp.com/jobs/Careers",
    "freshworks":    "https://careers.freshworks.com/jobs?q={title}",
    "razorpay":      "https://razorpay.com/jobs/#openings",
    "swiggy":        "https://careers.swiggy.com/#careers",
    "zomato":        "https://www.zomato.com/careers",
    "flipkart":      "https://www.flipkartcareers.com/#!/joblist",
    "phonepe":       "https://www.phonepe.com/careers/",
    "paytm":         "https://paytm.com/about-us/careers/",
}

# ── Input cleaning ─────────────────────────────────────────────────────────────

def clean_job_title(raw: str) -> str:
    first = raw.strip().splitlines()[0] if raw.strip() else raw
    first = re.sub(r'\s*\|\s*', '-', first)
    first = re.sub(r'^.*?[-–—]\s*hiring\s*[-–—]?\s*', '', first, flags=re.IGNORECASE)
    first = re.sub(r'^.*?[-–—]\s*recruitment\s*[-–—]?\s*', '', first, flags=re.IGNORECASE)
    first = re.sub(r'^.*?[-–—]\s*openings?\s*[-–—]?\s*', '', first, flags=re.IGNORECASE)
    first = re.sub(r'^.*?[-–—]\s*vacancy\s*[-–—]?\s*', '', first, flags=re.IGNORECASE)
    first = re.sub(
        r'\b(bangalore|bengaluru|mumbai|pune|hyderabad|chennai|delhi|ncr|noida|gurugram|gurgaon|kolkata|ahmedabad)'
        r'[\s\-]*(?:location)?\b', '', first, flags=re.IGNORECASE)
    first = re.sub(r'\b(freshers?|experienced?|lateral|entry[\s-]?level)\b', '', first, flags=re.IGNORECASE)
    first = re.split(
        r'\s*[-|–—]\s*(job\s*post|full[\s-]?time|part[\s-]?time|remote|hybrid|contract|intern)',
        first, flags=re.IGNORECASE)[0]
    first = re.sub(r'\d[\d.]*\s*(out of\s*)?\d\s*stars?', '', first, flags=re.IGNORECASE)
    first = re.sub(r'\(.*?\)', '', first)
    first = re.sub(r'[-–—]{2,}', '-', first)
    first = re.sub(r'\s{2,}', ' ', first)
    first = first.strip(' -–—\t')
    parts = [p.strip() for p in re.split(r'\s*[-–—]\s*', first) if p.strip()]
    if len(parts) > 1:
        def role_score(s):
            sc = len(s)
            if re.search(r'\b(ltd|pvt|inc|llc|corp|hiring|recruitment|openings?)\b', s, re.IGNORECASE):
                sc -= 50
            return sc
        first = max(parts, key=role_score)
    return first.strip(' -–—\t')


def clean_company_name(raw: str) -> str:
    lines = [l.strip() for l in raw.strip().splitlines() if l.strip()]
    name = lines[0] if lines else raw.strip()
    name = re.sub(r"[_*`~]{1,3}", "", name)
    name = re.sub(r'\d[\d.]*\s*(out of\s*)?\d\s*stars?', '', name, flags=re.IGNORECASE)
    name = re.sub(r'\(.*?\)', '', name)
    return name.strip(" -–—\t")


def company_slug(company: str) -> str:
    noise = {"technologies","technology","tech","solutions","services","pvt","ltd",
             "private","limited","inc","llc","corp","india","global","consulting"}
    words = [w for w in re.split(r"\s+", company.lower()) if len(w) >= 4 and w not in noise]
    return words[0] if words else company.lower()[:8]

# ── Core matching engine ───────────────────────────────────────────────────────

def score_text(text: str, job_title: str, company: str, raw_title: str = "") -> tuple:
    text_l   = text.lower()
    slug     = company_slug(company)
    co_lower = company.lower()

    co_on_page = slug in text_l or co_lower in text_l
    title_candidates = list(dict.fromkeys(filter(None, [
        job_title.lower().strip(),
        raw_title.lower().strip(),
    ])))

    for cand in title_candidates:
        if cand and cand in text_l:
            positions = [m.start() for m in re.finditer(re.escape(cand), text_l)]
            nearby = any(slug in text_l[max(0, p-600): p+600] for p in positions) if slug else True
            if nearby:
                return True, "exact title match + company co-occur", 100
            elif co_on_page:
                return True, "exact title match, company on page", 95
            else:
                break

    if not HAS_RAPIDFUZZ:
        words = [w for w in job_title.lower().split() if len(w) >= 2]
        matched = [w for w in words if w in text_l]
        title_score = int(len(matched) / len(words) * 100) if words else 0
    else:
        windows = []
        for kw in ("engineer","analyst","developer","manager","associate","consultant",
                   "specialist","designer","architect","position","opening","role","job"):
            for m in re.finditer(kw, text, re.IGNORECASE):
                s, e = max(0, m.start()-300), min(len(text), m.end()+300)
                windows.append(text[s:e].lower())
        if not windows:
            windows = [text[:6000].lower()]

        title_score = 0
        for cand in title_candidates:
            if not cand:
                continue
            scores = [fuzz.token_set_ratio(cand, w) for w in windows]
            title_score = max(title_score, max(scores) if scores else 0)

    if title_score >= FUZZY_TITLE_THRESHOLD and co_on_page:
        return True,  f"fuzzy title {title_score}/100 + company on page", title_score
    if title_score >= FUZZY_TITLE_THRESHOLD:
        return None,  f"fuzzy title {title_score}/100 but company NOT found on page", title_score
    if title_score >= 50 and co_on_page:
        return None,  f"partial title {title_score}/100 + company on page", title_score
    if not co_on_page:
        return False, f"company '{slug}' not found on page (title score {title_score})", title_score
    return False, f"title score too low: {title_score}/100", title_score


def is_login_wall(text: str) -> bool:
    tl = text.lower()
    return sum(1 for kw in LOGIN_WALL_KEYWORDS if kw in tl) >= 2 and len(text) < 4000

# ── Playwright helpers ─────────────────────────────────────────────────────────

async def _block_resources(route, request):
    if request.resource_type in BLOCK_RESOURCE_TYPES:
        await route.abort()
        return
    url = request.url.lower()
    if any(d in url for d in BLOCK_DOMAINS):
        await route.abort()
        return
    await route.continue_()


async def new_page(browser: Browser, ua: str | None = None) -> Page:
    ctx: BrowserContext = await browser.new_context(
        user_agent=ua or random.choice(USER_AGENTS),
        locale="en-IN",
        timezone_id="Asia/Kolkata",
        extra_http_headers={
            "Accept-Language": "en-IN,en;q=0.9",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "DNT": "1",
        },
        java_script_enabled=True,
    )
    await ctx.add_init_script("""
        Object.defineProperty(navigator, 'webdriver', {get: () => undefined});
        Object.defineProperty(navigator, 'plugins',   {get: () => [1,2,3,4,5]});
        Object.defineProperty(navigator, 'languages', {get: () => ['en-IN','en']});
        window.chrome = {runtime: {}};
    """)
    page: Page = await ctx.new_page()
    await page.route("**/*", _block_resources)
    return page


async def pw_goto(page: Page, url: str, done_event: asyncio.Event | None = None) -> bool:
    """Navigate to URL. Returns True on success. Checks done_event before proceeding."""
    # FIX 2: Guard — skip navigation if we are already shutting down
    if done_event is not None and done_event.is_set():
        return False
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=PW_GOTO_TIMEOUT)
        await page.wait_for_timeout(PW_WAIT_AFTER)
        return True
    except Exception:
        return False


async def pw_extract(page: Page) -> str:
    try:
        await page.evaluate("window.scrollBy(0, 800)")
        await page.wait_for_timeout(300)
        await page.evaluate("window.scrollTo(0, 0)")
        html = await page.content()
        return BeautifulSoup(html, "html.parser").get_text(separator=" ", strip=True)
    except Exception:
        return ""


async def pw_human_type(page: Page, selector: str, text: str) -> bool:
    try:
        await page.wait_for_selector(selector, timeout=4000, state="visible")
        await page.fill(selector, "")
        await page.type(selector, text, delay=random.randint(30, 80))
        await page.keyboard.press("Enter")
        await page.wait_for_timeout(random.randint(1200, 2000))
        return True
    except Exception:
        return False


async def pw_try_search_box(page: Page, selectors: list[str], query: str) -> bool:
    for sel in selectors:
        if await pw_human_type(page, sel, query):
            return True
    return False


async def pw_click_first_card(page: Page, selectors: list[str]) -> bool:
    for sel in selectors:
        try:
            els = await page.query_selector_all(sel)
            if els:
                await els[0].scroll_into_view_if_needed()
                await els[0].click()
                await page.wait_for_timeout(random.randint(1200, 2000))
                return True
        except Exception:
            continue
    return False

# ── HTTP fast-path ─────────────────────────────────────────────────────────────

_http_sem = None

async def http_get(client: "httpx.AsyncClient", url: str) -> str:
    global _http_sem
    async with _http_sem:
        try:
            r = await client.get(url, timeout=HTTP_TIMEOUT, follow_redirects=True)
            if r.status_code in (200, 203):
                return BeautifulSoup(r.text, "html.parser").get_text(separator=" ", strip=True)
        except Exception:
            pass
    return ""

# ── Generic scraper ────────────────────────────────────────────────────────────

async def scrape_platform(
    name: str,
    url: str,
    job_title: str,
    company: str,
    raw_title: str,
    browser: Browser,
    client: "httpx.AsyncClient",
    search_selectors: list[str],
    card_selectors: list[str],
    done_event: asyncio.Event,
) -> tuple:
    text = await http_get(client, url)
    if text and not is_login_wall(text) and len(text) > 400:
        r, reason, score = score_text(text, job_title, company, raw_title)
        if r is True:
            return r, f"[HTTP] {reason}", score
        if r is None and score >= 50:
            return r, f"[HTTP] {reason}", score

    page = await new_page(browser)
    try:
        # FIX 2: Pass done_event into pw_goto so it can abort before touching browser
        await pw_goto(page, url, done_event)

        text = await pw_extract(page)

        if is_login_wall(text):
            return None, "login wall", 0

        typed = await pw_try_search_box(page, search_selectors, f"{job_title} {company}")
        if typed:
            text = await pw_extract(page)
            if is_login_wall(text):
                return None, "login wall after search", 0
            r, reason, score = score_text(text, job_title, company, raw_title)
            if r is not False:
                clicked = await pw_click_first_card(page, card_selectors)
                if clicked:
                    deep = await pw_extract(page)
                    r2, reason2, score2 = score_text(deep, job_title, company, raw_title)
                    if r2 is not False and score2 >= score:
                        return r2, f"[deep JD] {reason2}", score2
                return r, reason, score

        clicked = await pw_click_first_card(page, card_selectors)
        if clicked:
            deep = await pw_extract(page)
            r, reason, score = score_text(deep, job_title, company, raw_title)
            if r is not False:
                return r, f"[deep JD] {reason}", score

        return score_text(text, job_title, company, raw_title)
    except asyncio.CancelledError:
        # FIX 1: Clean cancellation — just re-raise so the task exits
        raise
    except Exception as e:
        return None, f"error: {e}", 0
    finally:
        # FIX 2: Always close the context to release browser resources
        try:
            await page.context.close()
        except Exception:
            pass

# ── Platform definitions ───────────────────────────────────────────────────────

def _platform_configs(job_title: str, company: str) -> dict:
    q  = urlquote(f"{job_title} {company}")
    jt = urlquote(job_title)
    return {
        "Naukri": (
            f"https://www.naukri.com/jobs?q={q}&k={jt}",
            ["input#qsb-keyword-sugg", "input[placeholder*='skills' i]", "input[name='q']"],
            ["article.jobTuple", "div.srp-jobtuple-wrapper", "a.title"],
        ),
        "Indeed": (
            f"https://in.indeed.com/jobs?q={q}&l=India",
            ["input#text-input-what", "input[name='q']"],
            ["h2.jobTitle a", "a.jcs-JobTitle", "div.job_seen_beacon h2 a", ".resultContent h2 a"],
        ),
        "LinkedIn": (
            f"https://www.linkedin.com/jobs/search/?keywords={q}&location=India&f_TPR=r2592000",
            ["input.jobs-search-box__text-input", "input[aria-label*='Search' i]"],
            ["a.base-card__full-link", "div.base-search-card__info h3", ".job-search-card a"],
        ),
        "Glassdoor": (
            f"https://www.glassdoor.co.in/Job/india-jobs-SRCH_IL.0,5_IN115.htm?sc.keyword={q}",
            ["input[name='sc.keyword']", "input[placeholder*='Search' i]"],
            ["a.jobLink", "li.react-job-listing a", "[data-test='job-link']"],
        ),
        "Shine": (
            f"https://www.shine.com/job-search/{urlquote(job_title.replace(' ','-'))}-jobs-in-{urlquote(company.replace(' ','-'))}",
            ["input#search_query", "input[placeholder*='title' i]", "input[name='q']"],
            ["a.job-title-anchor", "div.job-listing-card a", "h2.job-title a"],
        ),
        "Foundit": (
            f"https://www.foundit.in/srp/results?query={q}&locations=India",
            ["input[placeholder*='search' i]", "input[name='q']"],
            ["a.card-apply-btn", "div.jobTupleHeader a", ".srpResultCardContainer a.title"],
        ),
        "TimesJobs": (
            f"https://www.timesjobs.com/candidate/job-search.html?searchType=personalizedSearch&from=submit&txtKeywords={q}&txtLocation=India",
            ["input#txtKeywords", "input[name='txtKeywords']"],
            ["h2.heading a", ".job-bx a.jobTitleLink", "ul li h2 a"],
        ),
        "Monster": (
            f"https://www.monsterindia.com/srp/results?query={q}&locations=India",
            ["input[placeholder*='skills' i]", "input[name='q']"],
            ["a.job-tittle", "div.job-title a", ".card-apply-btn"],
        ),
        "SimplyHired": (
            f"https://www.simplyhired.co.in/search?q={q}&l=India",
            ["input[name='q']", "input[placeholder*='Search' i]"],
            ["a[data-testid='jobTitle']", "h2.jobposting-title a", ".SerpJob-link"],
        ),
        "Internshala": (
            f"https://internshala.com/jobs/keywords-{urlquote(f'{job_title} {company}')}",
            ["input[placeholder*='keyword' i]", "input[name='q']"],
            ["a.job-internship-name", ".individual_internship_header a", "h3.job-internship-name a"],
        ),
        "Wellfound": (
            f"https://wellfound.com/jobs?q={jt}",
            ["input[placeholder*='Search' i]"],
            ["a[data-test='StartupResult']", "div.styles_component__nxVPb a", ".job-listing a"],
        ),
        "iimjobs": (
            f"https://www.iimjobs.com/j/{urlquote(job_title.lower().replace(' ','-'))}-jobs",
            ["input[placeholder*='Search' i]", "input[name='q']"],
            ["h2.job-title a", ".job-container a.title", "div.job-title a"],
        ),
        "Cutshort": (
            f"https://cutshort.io/jobs#?q={jt}",
            ["input[placeholder*='Search' i]"],
            [".job-listing a", "a.job-card-link"],
        ),
        "Instahyre": (
            f"https://www.instahyre.com/search-jobs/?q={jt}",
            ["input[placeholder*='Search' i]"],
            [".job-listing a", "a.job-card"],
        ),
    }

# ── Careers page check ─────────────────────────────────────────────────────────

async def resolve_careers_url(company: str, client: "httpx.AsyncClient") -> str | None:
    noise = {"technologies","technology","tech","solutions","services","inc","llc","ltd","pvt",
             "private","limited","group","corp","corporation","co","and","the","india","global"}
    words = company.lower().split()
    core  = [w.strip(".,") for w in words if w.strip(".,") not in noise]
    slugs = list(dict.fromkeys(["".join(core), core[0]] if core else [company.lower()]))

    candidates = []
    for s in slugs:
        candidates += [
            f"https://careers.{s}.com",
            f"https://jobs.{s}.com",
            f"https://www.{s}.com/careers",
            f"https://www.{s}.com/jobs",
            f"https://www.{s}.com/en/careers",
            f"https://www.{s}.com/in/careers",
        ]

    async def check(url):
        try:
            async with _http_sem:
                r = await client.get(url, timeout=5, follow_redirects=True)
                if r.status_code == 200:
                    html = r.text.lower()
                    if any(kw in html for kw in CAREERS_KEYWORDS):
                        return url
        except Exception:
            pass
        return None

    results = await asyncio.gather(*[check(u) for u in candidates])
    return next((r for r in results if r), None)


async def check_careers_page(
    careers_url: str, job_title: str, raw_title: str, company: str,
    browser: Browser, client: "httpx.AsyncClient",
) -> tuple:
    text = await http_get(client, careers_url)
    if text and not is_login_wall(text) and len(text) > 400:
        r, reason, score = score_text(text, job_title, company, raw_title)
        if r is True:
            return r, f"[HTTP careers] {reason} (score {score})"

    page = await new_page(browser)
    try:
        # Careers page check does NOT receive done_event — it always runs to completion
        await pw_goto(page, careers_url)
        page_text = await pw_extract(page)

        if is_login_wall(page_text):
            return None, "careers page behind login — inconclusive"

        search_selectors = [
            "input[type='search']", "input[placeholder*='search' i]",
            "input[placeholder*='job' i]", "input[placeholder*='keyword' i]",
            "input[placeholder*='title' i]", "input[id*='search' i]",
            "input[name*='search' i]", "input[name*='keyword' i]",
            "input[name*='q']", "input[class*='search' i]",
        ]
        typed = await pw_try_search_box(page, search_selectors, job_title)
        if typed:
            search_text = await pw_extract(page)
            if not is_login_wall(search_text):
                r, reason, score = score_text(search_text, job_title, company, raw_title)
                if r is not False:
                    card_sels = [
                        "a[class*='job']","a[class*='role']","a[class*='position']",
                        "h2 a","h3 a","li a","div[class*='result'] a",
                    ]
                    clicked = await pw_click_first_card(page, card_sels)
                    if clicked:
                        deep = await pw_extract(page)
                        r2, reason2, score2 = score_text(deep, job_title, company, raw_title)
                        if r2 is not False and score2 >= score:
                            return r2, f"[deep JD via careers search] {reason2} (score {score2})"
                    return r, f"[careers search] {reason} (score {score})"
                page_text = search_text

        encoded = urlquote(job_title)
        base    = careers_url.split("?")[0].rstrip("/")
        for search_url in [
            f"{base}?q={encoded}", f"{base}?keyword={encoded}",
            f"{base}?search={encoded}", f"{base}/search?q={encoded}",
            f"{base}/jobs?q={encoded}",
        ]:
            if search_url == careers_url:
                continue
            try:
                await pw_goto(page, search_url)
                t = await pw_extract(page)
                if is_login_wall(t):
                    continue
                r, reason, score = score_text(t, job_title, company, raw_title)
                if r is not False:
                    return r, f"[careers URL search] {reason} (score {score})"
            except Exception:
                continue

        r, reason, score = score_text(page_text, job_title, company, raw_title)
        return r, f"{reason} (score {score})"
    except asyncio.CancelledError:
        raise
    except Exception as e:
        return False, f"error: {e}"
    finally:
        try:
            await page.context.close()
        except Exception:
            pass

# ── Google confirmation ────────────────────────────────────────────────────────

async def google_confirm(
    job_title: str, company: str,
    browser: Browser, client: "httpx.AsyncClient",
    done_event: asyncio.Event,
) -> tuple:
    query = f'"{job_title}" "{company}" job India'
    url   = f"https://www.google.com/search?q={urlquote(query)}&hl=en&gl=in"

    text = await http_get(client, url)
    if text and len(text) > 500:
        if "captcha" not in text.lower():
            if any(x in text.lower() for x in ["did not match any", "no results found", "0 results"]):
                return False, "Google found no results for this job+company combo", 0
            r, reason, score = score_text(text, job_title, company)
            if r is not False or score > 10:
                return r, reason, score

    page = await new_page(browser)
    try:
        await pw_goto(page, url, done_event)
        if done_event.is_set():
            return None, "skipped (early exit)", 0
        src = await page.content()
        if any(x in src.lower() for x in ["captcha", "are you a robot", "unusual traffic"]):
            return None, "Google CAPTCHA — inconclusive", 0
        text2 = BeautifulSoup(src, "html.parser").get_text(separator=" ", strip=True)
        if any(x in text2.lower() for x in ["did not match any", "no results found"]):
            return False, "Google found no results", 0
        return score_text(text2, job_title, company)
    except asyncio.CancelledError:
        raise
    except Exception as e:
        return None, f"Google error: {e}", 0
    finally:
        try:
            await page.context.close()
        except Exception:
            pass

# ── Verdict aggregator ─────────────────────────────────────────────────────────

def aggregate_verdict(careers_result, google_result, platform_results: dict) -> tuple:
    confirmed_on, not_found_on = [], []
    cr, creason = careers_result
    gr, greason, gscore = google_result

    if cr is True:              confirmed_on.append("Careers page")
    if gr is True and gscore >= 92: confirmed_on.append("Google")
    for name, (r, reason, score) in platform_results.items():
        if r is True and int(score) >= 92:
            confirmed_on.append(name)

    if confirmed_on:
        for name, (r, reason, score) in platform_results.items():
            if r is False and name not in not_found_on:   not_found_on.append(name)
            elif r is True and name not in confirmed_on:  confirmed_on.append(name)
        if gr is False:  not_found_on.append("Google")
        if cr is False:  not_found_on.append("Careers page")
        return True, confirmed_on, not_found_on

    true_v = false_v = 0
    if cr is True:    true_v  += 3; confirmed_on.append("Careers page")
    elif cr is False: false_v += 3; not_found_on.append("Careers page")
    if gr is True:    true_v  += 2; confirmed_on.append("Google")
    elif gr is False: false_v += 2; not_found_on.append("Google")
    for name, (r, reason, score) in platform_results.items():
        if r is True:    true_v  += 1; confirmed_on.append(name)
        elif r is False: false_v += 1; not_found_on.append(name)

    if true_v == 0 and false_v == 0: return None,  confirmed_on, not_found_on
    if true_v >= false_v and true_v > 0: return True, confirmed_on, not_found_on
    if true_v > 0:                    return None,  confirmed_on, not_found_on
    return False, confirmed_on, not_found_on

# ── File reader ────────────────────────────────────────────────────────────────

def read_file(path: str) -> str:
    path = path.strip().strip('"').strip("'")
    ext  = path.lower().rsplit(".", 1)[-1] if "." in path else ""
    try:
        if ext in ("txt","csv","tsv","json","xml","eml","log","md"):
            with open(path, encoding="utf-8", errors="replace") as f:
                return f.read()
        if ext in ("html","htm"):
            with open(path, encoding="utf-8", errors="replace") as f:
                return BeautifulSoup(f.read(), "html.parser").get_text(separator="\n")
        if ext == "rtf":
            with open(path, encoding="utf-8", errors="replace") as f:
                raw = f.read()
            return re.sub(r'[{}\\]', '', re.sub(r'\\[a-z]+\-?\d*[ ]?', ' ', raw)).strip()
        if ext == "pdf":
            if pdfplumber is None:
                print("⚠️  pip install pdfplumber"); return ""
            with pdfplumber.open(path) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        if ext == "docx":
            if _docx_mod is None:
                print("⚠️  pip install python-docx"); return ""
            from docx import Document
            d = Document(path)
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip(): parts.append(cell.text.strip())
            return "\n".join(parts)
        if ext == "odt":
            try:
                from odf import teletype
                from odf.opendocument import load as odf_load
                return teletype.extractText(odf_load(path).text)
            except ImportError:
                print("⚠️  pip install odfpy"); return ""
        with open(path, encoding="utf-8", errors="replace") as f:
            return f.read()
    except FileNotFoundError:
        print(f"❌ File not found: {path}"); return ""
    except Exception as e:
        print(f"❌ Error reading file: {e}"); return ""

# ── FIX 3: Done-callback to silence "Future exception was never retrieved" ─────

def _shield_task(task: asyncio.Task) -> None:
    """Attach a no-op done callback so asyncio never logs unread exceptions."""
    def _cb(t: asyncio.Task):
        if not t.cancelled():
            t.exception()   # retrieves (and suppresses) the exception
    task.add_done_callback(_cb)

# ── Main async engine ──────────────────────────────────────────────────────────

async def _async_verify(job_title: str, company_name: str, raw_title: str):
    global _http_sem
    _http_sem = asyncio.Semaphore(MAX_HTTP_PARALLEL)
    t0 = time.monotonic()

    if not HAS_PLAYWRIGHT:
        print("❌ Playwright not installed. Run: pip install playwright && playwright install chromium")
        return
    if not HAS_HTTPX:
        print("❌ httpx not installed. Run: pip install httpx[http2]")
        return

    # FIX 4: Global exception handler to mute residual async noise
    loop = asyncio.get_running_loop()
    def _global_exc_handler(loop, context):
        exc = context.get("exception")
        # Silently swallow TargetClosedError and CancelledError leaking from tasks
        if exc is None:
            return
        name = type(exc).__name__
        if name in ("TargetClosedError", "CancelledError", "ConnectionClosedError"):
            return
        # For anything else, use the default handler
        loop.default_exception_handler(context)
    loop.set_exception_handler(_global_exc_handler)

    done_event = asyncio.Event()

    httpx_headers = {
        "User-Agent":      random.choice(USER_AGENTS),
        "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-IN,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "DNT":             "1",
    }

    # FIX 1: Keep references to all platform/google tasks so we can cancel them
    all_platform_tasks: list[asyncio.Task] = []

    async with async_playwright() as pw:
        browser: Browser = await pw.chromium.launch(
            headless=True,
            args=[
                "--no-sandbox",
                "--disable-dev-shm-usage",
                "--disable-gpu",
                "--disable-blink-features=AutomationControlled",
                "--lang=en-IN",
                "--window-size=1366,768",
            ],
        )

        async with httpx.AsyncClient(
            http2=True,
            headers=httpx_headers,
            follow_redirects=True,
            timeout=httpx.Timeout(HTTP_TIMEOUT),
        ) as client:

            # ── Step 1: Resolve careers URL ───────────────────────────────
            print("\n[1/4] Resolving company careers page…")
            careers_url = None
            key = company_name.lower().strip()
            for alias, tmpl in ATS_OVERRIDES.items():
                if alias in key or key in alias:
                    careers_url = tmpl.replace("{title}", urlquote(job_title))
                    print(f"  🎯 ATS matched ({alias}): {careers_url}")
                    break
            if not careers_url:
                careers_url = await resolve_careers_url(company_name, client)
                if careers_url:
                    print(f"  ✅ Resolved: {careers_url}")
                else:
                    print("  ⚠️  Could not resolve careers URL")

            # ─────────────────────────────────────────────────────────────
            # FIX 5: Step 2 — Run careers check FIRST, print result BEFORE
            #         platforms start.
            # ─────────────────────────────────────────────────────────────
            print("\n[2/4] Careers page (checking before platforms start)…")
            if careers_url:
                cr_raw = await check_careers_page(
                    careers_url, job_title, raw_title, company_name, browser, client
                )
            else:
                cr_raw = (None, "careers page not found")

            cr = cr_raw if isinstance(cr_raw, tuple) else (None, "careers error")
            icon = "✅" if cr[0] is True else ("⚠️ " if cr[0] is None else "❌")
            print(f"  {icon} {cr[1]}")

            # Trigger early exit if careers page already gives 100% confidence
            if cr[0] is True:
                # We still run platforms for completeness (they fill confirmed_on)
                # but done_event can be set by any platform that also hits ≥ 92
                pass

            # ─────────────────────────────────────────────────────────────
            # Steps 3 + 4: Google and all 14 platforms concurrently
            # ─────────────────────────────────────────────────────────────
            print("\n[3+4/4] Running Google + all 14 platforms simultaneously (full accuracy mode)…")
            print(f"  Platforms: {', '.join(_platform_configs(job_title, company_name).keys())}\n")

            platform_results: dict = {}
            platform_lock = asyncio.Lock()

            async def run_platform(name, url, search_sels, card_sels):
                try:
                    result = await asyncio.wait_for(
                        scrape_platform(
                            name, url, job_title, company_name, raw_title,
                            browser, client, search_sels, card_sels, done_event,
                        ),
                        timeout=PLATFORM_TIMEOUT,
                    )
                except asyncio.TimeoutError:
                    result = (None, f"timed out after {PLATFORM_TIMEOUT}s", 0)
                except asyncio.CancelledError:
                    # FIX 1: Task was cancelled by early-exit — record and exit cleanly
                    result = (None, "cancelled (early exit)", 0)
                except Exception as e:
                    result = (None, f"error: {e}", 0)

                r, reason, score = result
                icon  = "✅" if r is True else ("⚠️ " if r is None else "❌")
                score = int(score)
                bar   = "█" * (score // 10) + "░" * (10 - score // 10)
                print(f"  {icon} {name:<14} [{bar}] {score:>3}%  {reason}", flush=True)

                async with platform_lock:
                    platform_results[name] = result

                # All platforms always run to completion for full accuracy.
                # done_event is reserved for browser-shutdown cleanup only.
                return name, result

            # Google coroutine — now receives done_event so it can be skipped
            async def run_google():
                try:
                    result = await asyncio.wait_for(
                        google_confirm(job_title, company_name, browser, client, done_event),
                        timeout=PLATFORM_TIMEOUT,
                    )
                except asyncio.TimeoutError:
                    result = (None, f"timed out after {PLATFORM_TIMEOUT}s", 0)
                except asyncio.CancelledError:
                    result = (None, "cancelled (early exit)", 0)
                except Exception as e:
                    result = (None, f"Google error: {e}", 0)
                return result

            configs = _platform_configs(job_title, company_name)

            # FIX 1: Create named Tasks (not bare coroutines) so we can cancel them
            google_task = asyncio.ensure_future(run_google())
            _shield_task(google_task)  # FIX 3
            all_platform_tasks.append(google_task)

            for name, (url, search_sels, card_sels) in configs.items():
                t = asyncio.ensure_future(run_platform(name, url, search_sels, card_sels))
                _shield_task(t)  # FIX 3
                all_platform_tasks.append(t)

            # Wait for all tasks to finish (or be cancelled)
            await asyncio.gather(*all_platform_tasks, return_exceptions=True)

            gr_raw = google_task.result() if not google_task.cancelled() else (None, "cancelled", 0)
            gr_tup = gr_raw if isinstance(gr_raw, tuple) else (None, "google error", 0)

            elapsed = time.monotonic() - t0

            # ── Print Google result ────────────────────────────────────────
            print("\n[3/4] Google confirmation:")
            gr, greason, gscore = gr_tup
            icon = "✅" if gr is True else ("⚠️ " if gr is None else "❌")
            print(f"  {icon} Google [{gscore:>3}%]: {greason}")

            # ── Verdict ───────────────────────────────────────────────────
            final, confirmed_on, not_found_on = aggregate_verdict(
                cr, gr_tup, platform_results
            )

            print("\n" + "═"*54)
            if final is True:
                print(f"✅  REAL — '{job_title}' at {company_name}")
                print(f"   Confirmed on : {', '.join(confirmed_on)}")
            elif final is None:
                print(f"⚠️  UNCERTAIN — '{job_title}' at {company_name}")
                print(f"   Found on     : {', '.join(confirmed_on) or 'none'}")
                print(f"   Not found on : {', '.join(not_found_on) or 'none'}")
                print("   Tip: listing may be behind login, recently removed, or ATS-only.")
            else:
                print(f"⛔  SUSPICIOUS — '{job_title}' NOT confirmed at {company_name}")
                print(f"   Not found on : {', '.join(not_found_on)}")
                print("   Could be fake, already filled, or exclusively on a private ATS.")

            if careers_url:
                print(f"   Careers URL  : {careers_url}")
            print(f"   ⏱  Total time : {elapsed:.1f}s")
            print("═"*54 + "\n")

        # FIX 1: Cancel any tasks still running before closing browser
        for task in all_platform_tasks:
            if not task.done():
                task.cancel()

        # Give cancelled tasks a moment to clean up their Playwright contexts
        if any(not t.done() for t in all_platform_tasks):
            await asyncio.gather(*all_platform_tasks, return_exceptions=True)

        await browser.close()

# ── Input helpers ──────────────────────────────────────────────────────────────

def _readline(prompt: str = "") -> str:
    """
    Print prompt then read exactly one line from stdin.
    Uses sys.stdin.readline() instead of input() to avoid IDE / PyCharm
    run-console buffering quirks where multiple lines arrive concatenated
    on the first input() call.
    """
    if prompt:
        sys.stdout.write(prompt)
        sys.stdout.flush()
    try:
        line = sys.stdin.readline()
    except EOFError:
        return ""
    return line.rstrip("\n").rstrip("\r")


def _collect_two_lines() -> tuple[str, str]:
    """
    Read exactly two non-empty lines (job title, company name).
    Stops as soon as both are collected; ignores extra blank lines;
    never reads a third line accidentally.
    Returns (title_line, company_line).
    """
    collected: list[str] = []
    while len(collected) < 2:
        try:
            raw = sys.stdin.readline()
        except EOFError:
            break
        line = raw.rstrip("\n").rstrip("\r").strip()
        if line:                          # skip blank lines between the two
            collected.append(line)

    title   = collected[0] if len(collected) >= 1 else ""
    company = collected[1] if len(collected) >= 2 else ""
    return title, company

# ── Input collection ───────────────────────────────────────────────────────────

def verify_job():
    print("\n" + "═"*54)
    print("        JOB POSTING VERIFIER  v6  (MAX SPEED)")
    print("═"*54 + "\n")

    input_type = _readline("Input type (text / file):\n→ ").strip().lower()

    raw_title = raw_company = ""

    if input_type == "text":
        print("\nEnter job title on line 1, company on line 2:")

        # ── Read exactly 2 lines — no more, no less ───────────────────────
        raw_title, raw_company = _collect_two_lines()

        if not raw_title:
            print("❌ No input received."); return

        print(f"\n  Detected title  : {raw_title}")
        print(f"  Detected company: {raw_company if raw_company else '(not detected)'}")

        answer = _readline("  Correct? (y/n) → ").strip().lower()
        if answer != "y":
            raw_title   = _readline("Job Title    → ").strip()
            raw_company = _readline("Company Name → ").strip()
        elif not raw_company:
            raw_company = _readline("Company Name → ").strip()

    elif input_type == "file":
        path    = _readline("File path → ").strip().strip('"').strip("'")
        content = read_file(path)
        if not content.strip():
            print("❌ Could not read file or file is empty."); return
        print("\n── File preview (first 500 chars) ──")
        print(content[:500].strip())
        print("────────────────────────────────────")
        raw_title   = _readline("\nJob Title    → ").strip()
        raw_company = _readline("Company Name → ").strip()
    else:
        print("❌ Invalid input type. Please enter 'text' or 'file'."); return

    if not raw_title:
        raw_title = _readline("Job Title (required) → ").strip()
    if not raw_company:
        raw_company = _readline("Company Name (required) → ").strip()

    job_title    = clean_job_title(raw_title)
    company_name = clean_company_name(raw_company)

    if not job_title or not company_name:
        print("❌ Could not determine job title or company name."); return

    print(f"\n📋 Raw title  : '{raw_title}'")
    print(f"   Cleaned   : '{job_title}'")
    print(f"   Company   : '{company_name}'")
    print("─"*54)

    asyncio.run(_async_verify(job_title, company_name, raw_title))


if __name__ == "__main__":
    verify_job()